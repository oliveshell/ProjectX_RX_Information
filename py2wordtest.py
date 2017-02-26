# coding=utf-8
from docx import Document
from docx.shared import Inches
imgUrl = []
def initWord():
	document = Document()

	document.add_heading('Document Title', 0)

	p = document.add_paragraph('A plain paragraph having some ')
	p.add_run('bold').bold = True
	p.add_run(' and some ')
	p.add_run('italic.').italic = True

	document.add_heading('Heading, level 1', level=1)
	document.add_paragraph('Intense quote', style='IntenseQuote')




	import xlrd,os
	file = xlrd.open_workbook('./crawl_output_YS.xls')
	sheet= file.sheets()[0]
	nrows=sheet.nrows
	ncols = sheet.ncols
	
	for rownum in range(1,sheet.nrows):

		document.add_paragraph(sheet.cell(rownum,4).value, style='ListBullet')
		document.add_paragraph('     ')
		imgUrl.append(sheet.cell(rownum,5).value)
		path = 'img/'+str(rownum)+'.jpg'
		if os.path.exists(path):
			document.add_picture(path)
	

	#document.add_picture('monty-truth.png', width=Inches(1.25))

	
	# for item in recordset:
	#     row_cells = table.add_row().cells
	#     row_cells[0].text = str(item.qty)
	#     row_cells[1].text = str(item.id)
	#     row_cells[2].text = item.desc

	document.add_page_break()

	document.save('demo.docx')
def readImage():
	import urllib
	for url in imgUrl:
		if url:
			u = urllib.urlopen(url)
			data = u.read()
			fName = 'img/'+str(imgUrl.index(url))+'.jpg'
			f= open(fName,'wb')
			f.write(data)
			f.close()

def readExcel():

	import xlrd
	file = xlrd.open_workbook('./crawl_output_YS.xls')
	sheet= file.sheets()[0]
	nrows=sheet.nrows
	ncols = sheet.ncols
	for rownum in range(sheet.nrows):
		print sheet.cell(rownum,4).value
def createWord():
	initWord()
	readImage()
if __name__ =="__main__":
	createWord()